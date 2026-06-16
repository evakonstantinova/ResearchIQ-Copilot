from __future__ import annotations

import json
import os
import re
from typing import Any

import fitz  # PyMuPDF
import requests
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI

try:
    from docx import Document
except Exception:  # python-docx is optional until DOCX is uploaded
    Document = None

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("OPENAI_API_KEY not found. Create backend/.env with OPENAI_API_KEY=sk-...")

client = OpenAI(api_key=api_key)

app = FastAPI(title="ResearchIQ Copilot Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SEMANTIC_SCHOLAR_URL = "https://api.semanticscholar.org/graph/v1/paper/search"
MAX_FILE_CHARS = 35000
MAX_TOTAL_DOCUMENT_CHARS = 60000


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_pdf_text(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    if Document is None:
        raise HTTPException(
            status_code=500,
            detail="DOCX support requires python-docx. Install it with: pip install python-docx",
        )

    import io

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs).strip()


def extract_txt_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return ""


def trim_text(text: str, max_chars: int = MAX_FILE_CHARS) -> str:
    if len(text) <= max_chars:
        return text

    start = text[: int(max_chars * 0.45)]
    middle_start = max(0, len(text) // 2 - int(max_chars * 0.15))
    middle = text[middle_start : middle_start + int(max_chars * 0.25)]
    end = text[-int(max_chars * 0.30) :]

    return f"{start}\n\n[...middle section...]\n\n{middle}\n\n[...ending section...]\n\n{end}"


async def read_uploaded_file(file: UploadFile) -> dict[str, str]:
    filename = file.filename or "uploaded-file"
    lower = filename.lower()
    file_bytes = await file.read()

    if lower.endswith(".pdf"):
        text = extract_pdf_text(file_bytes)
    elif lower.endswith(".docx"):
        text = extract_docx_text(file_bytes)
    elif lower.endswith(".txt"):
        text = extract_txt_text(file_bytes)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")

    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"Could not extract text from {filename}. It may be scanned, empty, or unsupported.",
        )

    return {"filename": filename, "text": trim_text(text)}


def build_document_context(documents: list[dict[str, str]]) -> str:
    if not documents:
        return ""

    chunks: list[str] = []
    used = 0

    for index, doc in enumerate(documents, start=1):
        remaining = MAX_TOTAL_DOCUMENT_CHARS - used
        if remaining <= 0:
            break

        text = doc["text"][:remaining]
        used += len(text)
        chunks.append(
            f"--- ATTACHED DOCUMENT {index}: {doc['filename']} ---\n{text}\n--- END DOCUMENT {index} ---"
        )

    return "\n\n".join(chunks)


# ── Literature search ─────────────────────────────────────────────────────────

def looks_like_literature_request(message: str) -> bool:
    text = message.lower()
    triggers = [
        "literature",
        "similar paper",
        "similar papers",
        "related paper",
        "related papers",
        "related work",
        "find papers",
        "find articles",
        "find studies",
        "search papers",
        "search literature",
        "recent studies",
        "most cited",
        "references",
        "citation",
        "citations",
        "doi",
        "journal articles",
        "research gap",
        "systematic review",
    ]
    return any(trigger in text for trigger in triggers)


def create_literature_query(message: str, documents: list[dict[str, str]]) -> str:
    if not documents:
        return message[:280]

    doc_preview = "\n\n".join(
        f"Filename: {doc['filename']}\nContent preview:\n{doc['text'][:6000]}" for doc in documents[:2]
    )

    system = """
You create concise academic search queries for Semantic Scholar.
Return only one search query, no explanations.
Use the uploaded paper title, abstract, methods, datasets, and keywords when available.
Keep it under 18 words.
""".strip()

    user = f"User request: {message}\n\nAttached document preview:\n{doc_preview}"

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
            temperature=0.1,
        )
        query = response.choices[0].message.content.strip().strip('"')
        return query[:220] or message[:220]
    except Exception:
        return message[:220]


def search_semantic_scholar(query: str, limit: int = 10) -> list[dict[str, Any]]:
    params = {
        "query": query,
        "limit": min(max(limit, 1), 20),
        "fields": "title,authors,year,venue,journal,citationCount,externalIds,url,abstract,isOpenAccess,publicationTypes",
    }

    headers = {
        "User-Agent": "ResearchIQ-Copilot/1.0"
    }

    try:
        result = requests.get(
            SEMANTIC_SCHOLAR_URL,
            params=params,
            headers=headers,
            timeout=30,
        )
    except requests.RequestException as exc:
        print("Semantic Scholar connection error:", exc)
        return []

    if result.status_code != 200:
        print("Semantic Scholar status:", result.status_code)
        print("Semantic Scholar response:", result.text[:500])
        return []

    raw_papers = result.json().get("data", [])
    cleaned: list[dict[str, Any]] = []
    seen: set[str] = set()

    for paper in raw_papers:
        title = paper.get("title") or "Untitled"
        key = re.sub(r"\W+", "", title.lower())

        if not key or key in seen:
            continue

        seen.add(key)

        ext = paper.get("externalIds") or {}
        doi = ext.get("DOI")
        journal = paper.get("journal") or {}
        venue = paper.get("venue") or journal.get("name")

        cleaned.append(
            {
                "title": title,
                "authors": [
                    a.get("name")
                    for a in paper.get("authors", [])
                    if a.get("name")
                ],
                "year": paper.get("year"),
                "venue": venue,
                "citation_count": paper.get("citationCount"),
                "doi": doi,
                "doi_url": f"https://doi.org/{doi}" if doi else None,
                "url": paper.get("url"),
                "abstract": paper.get("abstract"),
                "is_open_access": paper.get("isOpenAccess"),
                "publication_types": paper.get("publicationTypes") or [],
            }
        )

    cleaned.sort(
        key=lambda p: ((p.get("year") or 0), (p.get("citation_count") or 0)),
        reverse=True,
    )

    return cleaned

def literature_context(papers: list[dict[str, Any]], query: str) -> str:
    if not papers:
        return f"Literature search query: {query}\nNo papers were returned."

    lines = [f"Literature search query: {query}", "Only use the papers listed below."]
    for i, paper in enumerate(papers, start=1):
        authors = ", ".join(paper.get("authors", [])[:5]) or "Not listed"
        abstract = paper.get("abstract") or "No abstract available."
        lines.append(
            f"""
[{i}] {paper.get('title')}
Authors: {authors}
Year: {paper.get('year') or 'n.d.'}
Venue: {paper.get('venue') or 'Not listed'}
Citations: {paper.get('citation_count') if paper.get('citation_count') is not None else 'Not listed'}
DOI: {paper.get('doi') or 'Not listed'}
URL: {paper.get('doi_url') or paper.get('url') or 'Not listed'}
Abstract: {abstract[:1200]}
""".strip()
        )
    return "\n\n".join(lines)


# ── OpenAI ────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """
You are ResearchIQ Copilot, a friendly expert AI research assistant.

You help users:
- chat with uploaded research documents
- summarise papers
- compare methods, datasets, findings, limitations, and research gaps
- search for real literature when literature results are provided
- write research-review style answers with clear structure

Rules:
- If uploaded documents are provided, use them directly and mention when something is not clearly stated.
- If literature search results are provided, only cite and discuss papers from those results.
- Never invent paper titles, authors, DOI links, citation counts, journals, metrics, or statistics.
- If live literature results are limited or imperfect, say so.
- Use clean markdown-style formatting. Use tables when the user asks for comparison.
- Be conversational but academically useful.
""".strip()


def call_openai(history: list[dict], user_message: str, context: str = "", temperature: float = 0.25) -> str:
    messages = [{"role": "system", "content": SYSTEM_PROMPT}]

    if context:
        messages.append(
            {
                "role": "system",
                "content": f"Context for this user message:\n\n{context}",
            }
        )

    clean_history = []
    for msg in history[-12:]:
        role = msg.get("role")
        content = msg.get("content")
        if role in {"user", "assistant"} and isinstance(content, str):
            clean_history.append({"role": role, "content": content[:6000]})

    messages.extend(clean_history)
    messages.append({"role": "user", "content": user_message})

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=temperature,
    )
    return response.choices[0].message.content


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/")
def home():
    return {"message": "ResearchIQ Copilot backend is running"}


@app.post("/chat")
async def chat(
    message: str = Form(""),
    history: str = Form("[]"),
    files: list[UploadFile] | None = File(default=None),
):
    user_message = message.strip() or "Please analyse the attached document."

    try:
        parsed_history = json.loads(history or "[]")
        if not isinstance(parsed_history, list):
            parsed_history = []
    except json.JSONDecodeError:
        parsed_history = []

    uploaded_files = files or []
    documents: list[dict[str, str]] = []

    for file in uploaded_files:
        documents.append(await read_uploaded_file(file))

    contexts: list[str] = []
    document_context = build_document_context(documents)
    if document_context:
        contexts.append(document_context)

    literature_results: list[dict[str, Any]] = []
    literature_query = ""

    if looks_like_literature_request(user_message):
        literature_query = create_literature_query(user_message, documents)
        literature_results = search_semantic_scholar(literature_query, limit=10)
        contexts.append(literature_context(literature_results, literature_query))

    if not document_context and not literature_results and looks_like_literature_request(user_message):
        contexts.append(
            "The user asked for literature, but the live search returned no usable papers. Be transparent and suggest better keywords."
        )

    context = "\n\n".join(contexts)

    answer = call_openai(
        history=parsed_history,
        user_message=user_message,
        context=context,
        temperature=0.22 if documents or literature_results else 0.38,
    )

    return {
        "answer": answer,
        "used_files": [doc["filename"] for doc in documents],
        "literature_search_used": bool(literature_results),
        "literature_query": literature_query,
        "literature_results_count": len(literature_results),
    }


@app.post("/find-literature")
async def find_literature_legacy(query: str = Form("")):
    """Optional legacy endpoint. The frontend no longer needs this."""
    q = query.strip()
    if not q:
        raise HTTPException(status_code=400, detail="Query is required.")
    return {"results": search_semantic_scholar(q, limit=10)}
